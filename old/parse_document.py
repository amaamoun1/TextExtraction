# -*- coding: utf-8 -*-
"""
This function reads in a path to a word document and uses docx to decompose the
word document into various parts, which are shipped to other functions for 
parsing.
At the moment, it reads in all of the tables in that word document and uses
which_table to only send the info and competency table to be parsed by
parse_table. This function then creates an overall document dictionary
connecting the parsed tables' names to each table dictionary. It also creates
a dataframe row with entries corresponding to the identified variables in all
of the tables.

To do:
    - Try to detect other relevant information
        - Table of Contents -- to get an overview of frequency of various 
                data in the overall dataset
        - Key Strengths and Key Risks
        - Firm-Specific Mission Text
        - Firm-Specific Key Outcomes Ratings/Comments
        - Company Specific Questions
        - Candidate Specific Recommendations
        - Career Goals and Motivations
    - Transform the entire document into one parseable text. Then check
        to see how many times certain words occur in the text using
        dictionaries or arrays of strings. This can be a crude linguistic
        approach i.e. measuring the frequency of words associated with
        overconfidence (overconfident, assertive, optimistic, etc.) versus
        the frequency of words associated with caution (cautious, timid,
        scientific, etc.).

Notes:
    - info on the docx package: https://python-docx.readthedocs.io/en/latest/index.html
    - info on the pandas package: https://pandas.pydata.org/
"""
import numpy as np
import pandas as pd
from docx import Document
from parse_table import parse_table
from which_table import which_table

def parse_document(doc_id, path, tables_to_parse):
    #set up document level data storage
    document = Document(path)
    doc_data = pd.DataFrame({'doc_id':[doc_id], 'doc_path':[path]})
    table_id = 0
    unparsed = 0
    table_tracker = pd.DataFrame(columns=['doc_id', 'table', 'unclean', 'parsed'])
    doc_unseen_keys=[]
    #loop through all tables in the document and parse those in tables_to_parse
    for table in document.tables:
        table_id += 1
        table_type = which_table(table)
        #keep track of the number of unparsed tables
        if table_type not in tables_to_parse:
            unparsed+=1
            continue
        #parse table and store info
#        print(table_type, "parsing")
        table_data, unclean_entries, table_unseen_keys = parse_table(table, table_type, doc_id) #parse the table and its dictionary and dataframe form
        table_tracker = table_tracker.append({'doc_id':doc_id, 'table': table_type, 'unclean entries': unclean_entries, 'parsed':1}, ignore_index=True, sort=True)
        doc_data = pd.merge(doc_data, table_data, on='doc_id', how='outer')
        doc_unseen_keys = doc_unseen_keys + table_unseen_keys
    #keep track of the tables_to_parse that were not found
    for tab in tables_to_parse:
        if sum(table_tracker["table"].str.contains(tab))==0:
            table_tracker = table_tracker.append({'doc_id':doc_id, 'table': tab, 'unclean': np.nan, 'parsed':0}, ignore_index=True, sort=True)
    table_tracker.set_index('table', 'doc_id', inplace=True)
    #if no info table found, print the first three paragraphs
#    if table_tracker.loc['info', 'parsed'] == 0:
#        for i in range(3):
#            p = document.paragraphs[i]
#            print(i, p.text)
    #lastly print number of unparsed tables and return info
    return (doc_data, table_tracker, doc_unseen_keys)












